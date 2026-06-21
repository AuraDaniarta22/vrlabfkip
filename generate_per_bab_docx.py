import re
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_document():
    doc = Document()
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(4.0)
        section.bottom_margin = Cm(3.0)
        section.left_margin = Cm(4.0)
        section.right_margin = Cm(3.0)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)
    return doc

def add_runs_with_italic(para, text, bold=False, font_size=12):
    parts = text.split('*')
    for i, part in enumerate(parts):
        if part == '':
            continue
        run = para.add_run(part)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(font_size)
        run.font.bold = bold
        if i % 2 == 1:
            run.font.italic = True

def process_file(doc, filepath):
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()
    lines = content.split('\n')
    
    heading_patterns = [
        r'^Kajian tentang', r'^Laboratorium Fisik', r'^Penelitian Terdahulu',
        r'^Kerangka Berpikir', r'^Latar Belakang', r'^Rumusan Masalah',
        r'^Tujuan Penelitian', r'^Batasan Masalah', r'^Manfaat Penelitian',
        r'^Teoritis$', r'^Praktis$', r'^Bagi ', r'^Populasi$', r'^Sampel$',
        r'^Subjek Uji', r'^Prosedur Penelitian', r'^Tahap ', r'^Analisis ',
        r'^Perancangan ', r'^Desain Responsif', r'^Pengambilan dan Optimasi',
        r'^Proses ', r'^Implementasi', r'^Instrumen ', r'^Angket ',
        r'^Teknik ', r'^Evaluasi', r'^Tujuan Implementasi',
        r'^Jenis dan Pendekatan', r'^Populasi dan Sampel',
        r'^Ahli Materi$', r'^Ahli Media$', r'^Pendahuluan$',
        r'^Kebutuhan Sistem$', r'^Cara Mengakses', r'^Panduan Navigasi',
        r'^Daftar Area', r'^Tips Tambahan', r'^Navigasi ', r'^Peta Interaktif',
        r'^Tombol Kontrol', r'^Halaman ', r'^Area ', r'^Elemen ',
        r'^Konversi ke', r'^Pembuatan ', r'^Konfigurasi ', r'^Penanganan ',
        r'^Integrasi Google', r'^Kompatibilitas ', r'^Fungsionalitas ',
        r'^Kinerja ', r'^Kualitas Visual', r'^Teknologi ', r'^Struktur Hierarki',
        r'^Analisis Data', r'^3DVista',
    ]

    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue

        # BAB headings
        if line.startswith('BAB '):
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(12)
            para.paragraph_format.line_spacing = 1.5
            para.paragraph_format.first_line_indent = None
            run = para.add_run(line)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.font.bold = True
            i += 1
            continue

        # Sub-chapter numbering (4.1, 4.1.1, 5.1, 5.2)
        if re.match(r'^\d+\.\d+', line):
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(6)
            para.paragraph_format.line_spacing = 1.5
            para.paragraph_format.first_line_indent = None
            add_runs_with_italic(para, line, bold=True)
            i += 1
            continue

        # A. B. C. D. headings
        if re.match(r'^[A-Z]\.\s', line):
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)
            para.paragraph_format.line_spacing = 1.5
            para.paragraph_format.first_line_indent = None
            add_runs_with_italic(para, line, bold=True)
            i += 1
            continue

        # Check heading patterns
        clean_line = line.replace('*', '')
        is_heading = False
        for pattern in heading_patterns:
            if re.match(pattern, clean_line):
                is_heading = True
                break
        
        if is_heading and len(line) < 120:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(3)
            para.paragraph_format.line_spacing = 1.5
            para.paragraph_format.first_line_indent = None
            add_runs_with_italic(para, line, bold=True)
            i += 1
            continue

        # Tabel references
        if line.startswith('Tabel ') and len(line) < 120:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)
            para.paragraph_format.line_spacing = 1.5
            para.paragraph_format.first_line_indent = None
            add_runs_with_italic(para, line, bold=True)
            i += 1
            continue

        # Gambar references
        if line.startswith('Gambar ') and len(line) < 120:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)
            para.paragraph_format.line_spacing = 1.5
            para.paragraph_format.first_line_indent = None
            run = para.add_run('[TEMPATKAN ' + line.upper() + ' DI SINI]')
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.font.italic = True
            i += 1
            continue

        # Sumber references
        if line.startswith('Sumber:'):
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.space_before = Pt(3)
            para.paragraph_format.space_after = Pt(6)
            para.paragraph_format.line_spacing = 1.5
            para.paragraph_format.first_line_indent = None
            add_runs_with_italic(para, line, font_size=10)
            i += 1
            continue

        # Numbered list
        if re.match(r'^\d+[\.\)]\s', line.strip()):
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.space_after = Pt(3)
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.line_spacing = 1.5
            para.paragraph_format.left_indent = Cm(1.27)
            para.paragraph_format.first_line_indent = None
            add_runs_with_italic(para, line)
            i += 1
            continue

        # Regular paragraph
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.space_after = Pt(6)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.line_spacing = 1.5
        para.paragraph_format.first_line_indent = Cm(1.27)
        add_runs_with_italic(para, line)
        i += 1

def main():
    bab_files = {
        1: r'c:\skripsi\vrlabfkip\bab1_revisi.txt',
        2: r'c:\skripsi\vrlabfkip\bab2_revisi.txt',
        3: r'c:\skripsi\vrlabfkip\bab3_revisi.txt',
        4: r'c:\skripsi\vrlabfkip\bab4_revisi.txt',
        5: r'c:\skripsi\vrlabfkip\bab5_revisi.txt',
    }
    
    for bab_num, bab_file in bab_files.items():
        doc = create_document()
        process_file(doc, bab_file)
        output_path = rf'c:\skripsi\vrlabfkip\Revisi_BAB_{bab_num}.docx'
        doc.save(output_path)
        print(f"Saved: {output_path}")
    
    print("\nSemua file .docx per bab berhasil dibuat!")

if __name__ == '__main__':
    main()
