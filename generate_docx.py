import re
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def create_document():
    doc = Document()
    
    # Page setup: A4, margins
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(4.0)
        section.bottom_margin = Cm(3.0)
        section.left_margin = Cm(4.0)
        section.right_margin = Cm(3.0)
    
    # Default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)
    
    return doc

def add_paragraph_with_italic(doc, text, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False, font_size=12, space_after=Pt(6), first_line_indent=Cm(1.27)):
    """Add a paragraph with italic support for text between * markers."""
    para = doc.add_paragraph()
    para.alignment = alignment
    para.paragraph_format.space_after = space_after
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.line_spacing = 1.5
    if first_line_indent:
        para.paragraph_format.first_line_indent = first_line_indent
    
    # Split text by * for italic handling
    parts = text.split('*')
    for i, part in enumerate(parts):
        if part == '':
            continue
        run = para.add_run(part)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(font_size)
        run.font.bold = bold
        if i % 2 == 1:  # odd index = italic
            run.font.italic = True
    
    return para

def process_file(doc, filepath):
    """Read a text file and add its content to the document."""
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # Skip empty lines
        if not line:
            i += 1
            continue
        
        # BAB headings (centered, bold, large)
        if line.startswith('BAB '):
            # Check if next non-empty line is part of the title
            title_lines = [line]
            j = i + 1
            while j < len(lines) and lines[j].strip() == '':
                j += 1
            
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(12)
            para.paragraph_format.line_spacing = 1.5
            para.paragraph_format.first_line_indent = None
            run = para.add_run(line)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.font.bold = True
            i += 1
            continue
        
        # Sub-chapter numbering patterns like "4.1", "4.1.1", "4.2.1"
        if re.match(r'^\d+\.\d+', line):
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(6)
            para.paragraph_format.line_spacing = 1.5
            para.paragraph_format.first_line_indent = None
            
            # Handle italic in heading
            parts = line.split('*')
            for idx, part in enumerate(parts):
                if part == '':
                    continue
                run = para.add_run(part)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.font.bold = True
                if idx % 2 == 1:
                    run.font.italic = True
            i += 1
            continue
        
        # Section headings (lines starting with A., B., C., etc.)
        if re.match(r'^[A-Z]\.\s', line):
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)
            para.paragraph_format.line_spacing = 1.5
            para.paragraph_format.first_line_indent = None
            
            parts = line.split('*')
            for idx, part in enumerate(parts):
                if part == '':
                    continue
                run = para.add_run(part)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.font.bold = True
                if idx % 2 == 1:
                    run.font.italic = True
            i += 1
            continue
        
        # Sub-section headings (lines that are short and look like titles)
        # Detect headings: lines with no period at end, relatively short, and common heading patterns
        heading_patterns = [
            r'^Kajian tentang',
            r'^Laboratorium Fisik',
            r'^Penelitian Terdahulu',
            r'^Kerangka Berpikir',
            r'^Latar Belakang',
            r'^Rumusan Masalah',
            r'^Tujuan Penelitian',
            r'^Batasan Masalah',
            r'^Manfaat Penelitian',
            r'^Teoritis$',
            r'^Praktis$',
            r'^Bagi ',
            r'^Populasi$',
            r'^Sampel$',
            r'^Subjek Uji',
            r'^Prosedur Penelitian',
            r'^Tahap ',
            r'^Analisis ',
            r'^Perancangan ',
            r'^Desain Responsif',
            r'^Pengambilan dan Optimasi',
            r'^Proses ',
            r'^Implementasi',
            r'^Instrumen ',
            r'^Angket ',
            r'^Teknik ',
            r'^Evaluasi',
            r'^Tujuan Implementasi',
            r'^Jenis dan Pendekatan',
            r'^Populasi dan Sampel',
            r'^Ahli Materi$',
            r'^Ahli Media$',
            r'^Pendahuluan$',
            r'^Kebutuhan Sistem$',
            r'^Cara Mengakses',
            r'^Panduan Navigasi',
            r'^Daftar Area',
            r'^Tips Tambahan',
            r'^Navigasi ',
            r'^Peta Interaktif',
            r'^Tombol Kontrol',
            r'^Halaman ',
            r'^Area ',
            r'^Elemen ',
            r'^Konversi ke',
            r'^Pembuatan ',
            r'^Konfigurasi ',
            r'^Penanganan ',
            r'^Integrasi Google',
            r'^Kompatibilitas ',
            r'^Fungsionalitas ',
            r'^Kinerja ',
            r'^Kualitas Visual',
            r'^Teknologi ',
            r'^Kebutuhan Sistem',
            r'^Struktur Hierarki',
        ]
        
        is_heading = False
        for pattern in heading_patterns:
            if re.match(pattern, line.replace('*', '')):
                is_heading = True
                break
        
        # Also detect lines starting with "Level" for sub-items
        if re.match(r'^\*Level\*', line) or re.match(r'^Level', line.replace('*', '')):
            is_heading = False  # These are descriptions, not headings
        
        if is_heading and len(line) < 120:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(3)
            para.paragraph_format.line_spacing = 1.5
            para.paragraph_format.first_line_indent = None
            
            parts = line.split('*')
            for idx, part in enumerate(parts):
                if part == '':
                    continue
                run = para.add_run(part)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.font.bold = True
                if idx % 2 == 1:
                    run.font.italic = True
            i += 1
            continue
        
        # Table references
        if line.startswith('Tabel ') and len(line) < 120:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)
            para.paragraph_format.line_spacing = 1.5
            para.paragraph_format.first_line_indent = None
            
            parts = line.split('*')
            for idx, part in enumerate(parts):
                if part == '':
                    continue
                run = para.add_run(part)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.font.bold = True
                if idx % 2 == 1:
                    run.font.italic = True
            i += 1
            continue
        
        # Gambar references
        if line.startswith('Gambar ') and len(line) < 120:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)
            para.paragraph_format.line_spacing = 1.5
            para.paragraph_format.first_line_indent = None
            run = para.add_run('[TEMPATKAN ' + line.upper() + ' DI SINI]')
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.font.italic = True
            run.font.color.rgb = None
            i += 1
            continue
        
        # Sumber references
        if line.startswith('Sumber:'):
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.space_before = Pt(3)
            para.paragraph_format.space_after = Pt(6)
            para.paragraph_format.line_spacing = 1.5
            para.paragraph_format.first_line_indent = None
            
            parts = line.split('*')
            for idx, part in enumerate(parts):
                if part == '':
                    continue
                run = para.add_run(part)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
                if idx % 2 == 1:
                    run.font.italic = True
            i += 1
            continue
        
        # Numbered list items (1. 2. 3. etc.)
        if re.match(r'^\d+[\.\)]\s', line) or line.startswith('\t') and re.match(r'^\d+[\.\)]\s', line.strip()):
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.space_after = Pt(3)
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.line_spacing = 1.5
            para.paragraph_format.left_indent = Cm(1.27)
            para.paragraph_format.first_line_indent = None
            
            parts = line.split('*')
            for idx, part in enumerate(parts):
                if part == '':
                    continue
                run = para.add_run(part)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                if idx % 2 == 1:
                    run.font.italic = True
            i += 1
            continue
        
        # Regular paragraph
        add_paragraph_with_italic(doc, line)
        i += 1

def main():
    doc = create_document()
    
    # Process each chapter
    bab_files = [
        r'c:\skripsi\vrlabfkip\bab1_revisi.txt',
        r'c:\skripsi\vrlabfkip\bab2_revisi.txt',
        r'c:\skripsi\vrlabfkip\bab3_revisi.txt',
        r'c:\skripsi\vrlabfkip\bab4_revisi.txt',
        r'c:\skripsi\vrlabfkip\bab5_revisi.txt',
    ]
    
    for bab_file in bab_files:
        process_file(doc, bab_file)
        # Add page break after each chapter except the last
        if bab_file != bab_files[-1]:
            doc.add_page_break()
    
    # Add Daftar Pustaka
    doc.add_page_break()
    
    # Read the daftar pustaka from original
    with open(r'c:\skripsi\vrlabfkip\skripsi.txt', 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Find DAFTAR PUSTAKA section
    dp_start = content.find('DAFTAR PUSTAKA')
    dp_end = content.find('Lampiran 1.')
    if dp_start >= 0 and dp_end >= 0:
        dp_text = content[dp_start:dp_end].strip()
    elif dp_start >= 0:
        dp_text = content[dp_start:].strip()
    else:
        dp_text = "DAFTAR PUSTAKA"
    
    dp_lines = dp_text.split('\n')
    for line in dp_lines:
        line = line.strip()
        if not line:
            continue
        if line == 'DAFTAR PUSTAKA':
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(12)
            para.paragraph_format.line_spacing = 1.5
            para.paragraph_format.first_line_indent = None
            run = para.add_run(line)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.font.bold = True
        else:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.space_after = Pt(6)
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.line_spacing = 1.5
            para.paragraph_format.left_indent = Cm(1.27)
            para.paragraph_format.first_line_indent = Cm(-1.27)
            run = para.add_run(line)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
    
    output_path = r'c:\skripsi\vrlabfkip\Skripsi_Revisi_Aura_Daniarta_2025.docx'
    doc.save(output_path)
    print(f"Document saved to: {output_path}")

if __name__ == '__main__':
    main()
